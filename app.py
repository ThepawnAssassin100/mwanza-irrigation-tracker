import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime

# Function to create and export Word document
def to_word(filtered, officer, epa, ta):
    doc = Document()
    doc.add_heading('Mwanza District Irrigation Task Report 🌱📋', 0)

    doc.add_paragraph(f"🧑‍💼 Responsible Officer: {officer}")
    doc.add_paragraph(f"📍 EPA: {epa if epa else 'N/A'}")
    doc.add_paragraph(f"🏘️ T/A: {ta if ta else 'N/A'}")
    doc.add_paragraph(f"📅 Date: {datetime.now().strftime('%Y-%m-%d')}")

    doc.add_heading('📝 Task Details:', level=1)
    table = doc.add_table(rows=1, cols=len(filtered.columns))

    # Adding headers
    for i, col in enumerate(filtered.columns):
        table.cell(0, i).text = col

    # Adding rows
    for _, row in filtered.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

# Function to export to Excel
def to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Tasks')
    output.seek(0)
    return output

# Streamlit interface
def main():
    st.set_page_config(page_title="Mwanza Irrigation Tracker", page_icon="🌾", layout="wide")
    st.markdown("""
        <h1 style="color:#4CAF50; font-size:40px; text-align:center;">🌾 Mwanza District Irrigation Tracker 🌦️</h1>
        <p style="font-size:18px; text-align:center;">📌 Track field and office tasks 🛠️ with responsible officers, descriptions, remarks, and location info. Export monthly reports in Word or Excel. 📤</p>
    """, unsafe_allow_html=True)

    # Initial task data
    if 'tasks' not in st.session_state:
        st.session_state.tasks = pd.DataFrame(columns=[
            '📌 Task', '📝 Work Description', '📆 Deadline', '📊 Status', '🚧 Progress', 
            '🏢 Office Work', '💬 Remarks', '🗓️ Created On'
        ])

    st.subheader("➕ Add New Task")
    task = st.text_input("📌 Task")
    description = st.text_input("📝 Work Description")
    deadline = st.date_input("📆 Deadline")
    status = st.selectbox("📊 Status", ["In Progress 🚧", "Completed ✅", "Pending ⏳"])
    progress = st.text_input("🚧 Progress (e.g., 50%)")
    office_work = st.text_input("🏢 Office Work")
    remarks = st.text_input("💬 Remarks")
    photo = st.file_uploader("📸 Upload Photo (Optional)", type=["jpg", "jpeg", "png"])
    if st.button("📥 Add Task"):
        if task:
            new_row = pd.DataFrame.from_dict([{
                '📌 Task': task,
                '📝 Work Description': description,
                '📆 Deadline': deadline.strftime('%Y-%m-%d'),
                '📊 Status': status,
                '🚧 Progress': progress,
                '🏢 Office Work': office_work,
                '💬 Remarks': remarks,
                '🗓️ Created On': datetime.today().strftime('%Y-%m-%d')
            }])
            st.session_state.tasks = pd.concat([st.session_state.tasks, new_row], ignore_index=True)

    st.subheader("📋 Current Tasks")
    if not st.session_state.tasks.empty:
        edited_df = st.experimental_data_editor(st.session_state.tasks, num_rows="dynamic")
        st.session_state.tasks = edited_df

    officer = st.text_input("🧑‍💼 Responsible Officer")
    epa = st.text_input("📍 EPA (Optional)")
    ta = st.text_input("🏘️ T/A (Optional)")

    selected_month = st.date_input("📅 Select Month", datetime.today())
    month_str = selected_month.strftime("%Y-%m")
    filtered_df = st.session_state.tasks[st.session_state.tasks['📆 Deadline'].str.startswith(month_str)]

    st.write(f"📅 Filtered tasks for {selected_month.strftime('%B %Y')}:")
    st.dataframe(filtered_df)

    if st.button('📄 Download Word Report'):
        word_data = to_word(filtered_df, officer, epa, ta)
        st.download_button(
            label="📥 Download Word Document",
            data=word_data,
            file_name=f'mwanza_irrigation_report_{month_str}.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    if st.button("📊 Download Excel Report"):
        excel_data = to_excel(filtered_df)
        st.download_button(
            label="📥 Download Excel File",
            data=excel_data,
            file_name=f"irrigation_tasks_{month_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

if __name__ == '__main__':
    main()
